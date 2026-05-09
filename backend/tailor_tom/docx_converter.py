"""DOCX to LaTeX conversion using python-docx and OpenAI.

Pipeline (Option D):
  1. python-docx extracts structured content from the .docx file.
  2. An LLM classifies that content into a structured JSON schema
     (section types, entry fields) — one cheap call.
  3. A deterministic Python renderer converts the JSON into LaTeX
     with pre-tuned templates.
  4. A binary-search page-fitting loop adjusts spacing and recompiles
     locally (no extra LLM calls).
"""

import io
import json
import logging
import re
import traceback
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from openai import OpenAI

from tailor_tom.config import settings
from tailor_tom.resume_renderer import fit_to_pages

logger = logging.getLogger(__name__)


class DOCXConversionError(RuntimeError):
    """DOCX conversion failure with admin-facing debug context."""

    def __init__(self, message: str, debug_context: dict[str, Any]):
        super().__init__(message)
        self.debug_context = debug_context


def _debug_preview(value: str, max_len: int = 2000) -> str:
    """Compact preview for admin diagnostics without storing unbounded LLM output."""
    compact = " ".join((value or "").split())
    if len(compact) <= max_len:
        return compact
    return compact[: max_len - 3] + "..."


def _usage_to_debug_dict(usage: Any) -> dict[str, Any] | None:
    """Best-effort OpenAI usage serialization for JSON debug payloads."""
    if usage is None:
        return None
    if hasattr(usage, "model_dump"):
        try:
            return usage.model_dump()
        except Exception:
            pass
    if isinstance(usage, dict):
        return usage
    return {"repr": repr(usage)}


# ---------------------------------------------------------------------------
# Step 1: Extract structured content from .docx
# ---------------------------------------------------------------------------


def extract_content_from_docx(docx_bytes: bytes) -> dict:
    """Parse a .docx file and extract structured content with formatting metadata.

    Args:
        docx_bytes: Raw bytes of the .docx file.

    Returns:
        A dict with keys:
          - elements: list of element dicts (paragraph/table) in document order
          - raw_text: plain-text concatenation of all content
    """
    doc = Document(io.BytesIO(docx_bytes))

    elements: list[dict] = []
    raw_text_parts: list[str] = []

    # Process paragraphs and tables in document order
    for block in _iter_block_items(doc):
        if block["type"] == "paragraph":
            para = block["element"]
            para_data = _extract_paragraph(para)
            if para_data:
                elements.append(para_data)
                raw_text_parts.append(para_data["text"])

        elif block["type"] == "table":
            table = block["element"]
            table_data = _extract_table(table)
            if table_data:
                elements.append(table_data)
                for row in table_data["rows"]:
                    for cell in row:
                        raw_text_parts.append(cell["text"])

    return {
        "elements": elements,
        "raw_text": "\n".join(raw_text_parts),
    }


def _iter_block_items(doc: Document) -> list[dict]:
    """Iterate over paragraphs and tables in document body order."""
    items = []
    body = doc.element.body
    for child in body:
        if child.tag == qn("w:p"):
            # Find the matching paragraph object
            for para in doc.paragraphs:
                if para._element is child:
                    items.append({"type": "paragraph", "element": para})
                    break
        elif child.tag == qn("w:tbl"):
            for table in doc.tables:
                if table._element is child:
                    items.append({"type": "table", "element": table})
                    break
    return items


def _extract_paragraph(para) -> Optional[dict]:
    """Extract a single paragraph with formatting metadata, including tab stops."""
    text = para.text.strip()
    if not text:
        return None

    # Determine alignment
    alignment = "left"
    if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        alignment = "center"
    elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        alignment = "right"

    # Detect style (heading, list item, normal)
    style_name = para.style.name if para.style else "Normal"
    is_heading = style_name.startswith("Heading")
    heading_level = 0
    if is_heading:
        try:
            heading_level = int(style_name.replace("Heading", "").strip())
        except ValueError:
            heading_level = 1

    is_list = style_name.startswith("List") or "bullet" in style_name.lower()

    # Extract tab stop definitions from paragraph format (right/center/left tabs)
    tab_stop_type = _get_tab_stop_alignment(para)

    # Extract runs with formatting, detecting tab characters from XML
    runs = []
    has_tabs = False
    for run_element in para._element.findall(qn("w:r")):
        # Check for tab characters (<w:tab/>) inside this run
        for child in run_element:
            if child.tag == qn("w:tab"):
                has_tabs = True
                runs.append({"text": "\t", "tab": True})
            elif child.tag == qn("w:t"):
                run_text = child.text or ""
                if not run_text:
                    continue
                # Get formatting from the run's properties
                run_obj = None
                for r in para.runs:
                    if r._element is run_element:
                        run_obj = r
                        break
                run_data: dict = {"text": run_text}
                if run_obj:
                    if run_obj.bold:
                        run_data["bold"] = True
                    if run_obj.italic:
                        run_data["italic"] = True
                    if run_obj.underline:
                        run_data["underline"] = True
                    if run_obj.font.size:
                        run_data["font_size_pt"] = run_obj.font.size.pt
                    if run_obj.font.name:
                        run_data["font_name"] = run_obj.font.name
                runs.append(run_data)

    result = {
        "type": "paragraph",
        "text": text,
        "alignment": alignment,
        "style": style_name,
        "is_heading": is_heading,
        "heading_level": heading_level,
        "is_list": is_list,
        "runs": runs,
    }

    if has_tabs and tab_stop_type:
        result["tab_alignment"] = tab_stop_type

    return result


def _get_tab_stop_alignment(para) -> Optional[str]:
    """Get the alignment type of the first tab stop defined on a paragraph.

    Returns 'right', 'center', or None if no tab stops defined.
    """
    pPr = para._element.find(qn("w:pPr"))
    if pPr is None:
        return None
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        return None
    for tab in tabs.findall(qn("w:tab")):
        val = tab.get(qn("w:val"))
        if val in ("right", "center"):
            return val
    return None


def _extract_table(table) -> Optional[dict]:
    """Extract a table with cell content and formatting."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            # Get alignment from the first paragraph in the cell
            alignment = "left"
            if cell.paragraphs:
                first_para = cell.paragraphs[0]
                if first_para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment = "center"
                elif first_para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment = "right"
            cells.append({"text": cell_text, "alignment": alignment})
        rows.append(cells)

    # Skip empty tables
    if not any(cell["text"] for row in rows for cell in row):
        return None

    return {
        "type": "table",
        "rows": rows,
        "num_rows": len(rows),
        "num_cols": len(rows[0]) if rows else 0,
    }


# ---------------------------------------------------------------------------
# Marker sanitization (defense in depth: strip conversion artifacts from JSON)
# ---------------------------------------------------------------------------

# Patterns that must never appear in final output (bracket/alignment tokens from LLM input)
_DOCX_MARKER_PATTERNS: list[tuple[re.Pattern[str], str]] = [
    # Combined modifier tags (any order)
    (re.compile(r"\[(?:BOLD|ITALIC|UNDERLINE)(?:,(?:BOLD|ITALIC|UNDERLINE))*\]", re.I), ""),
    (re.compile(r"\[/\]"), ""),
    (re.compile(r"\[TAB\]"), ""),
    (re.compile(r"\[RIGHT-ALIGNED\]"), ""),
    (re.compile(r"\[CENTER-ALIGNED\]"), ""),
    (re.compile(r"\(align:\s*center\)", re.I), ""),
    (re.compile(r"\(align:\s*right\)", re.I), ""),
    (re.compile(r"\(align:\s*left\)", re.I), ""),
    (re.compile(r"\[HEADING\s+level=\d+\]\s*", re.I), ""),
    (re.compile(r"\[BULLET\]\s*", re.I), ""),
    (re.compile(r"\[TABLE\s+\d+x\d+\]", re.I), ""),
    (re.compile(r"\[/TABLE\]", re.I), ""),
]

# Single pattern to detect any of the above (for safety gate scan)
_DOCX_MARKER_ANY = re.compile(
    r"\[(?:BOLD|ITALIC|UNDERLINE)(?:,(?:BOLD|ITALIC|UNDERLINE))*\]|"
    r"\[/\]|\[TAB\]|\[RIGHT-ALIGNED\]|\[CENTER-ALIGNED\]|"
    r"\(align:\s*(?:center|right|left)\)|"
    r"\[HEADING\s+level=\d+\]|\[BULLET\]|\[TABLE\s+\d+x\d+\]|\[/TABLE\]",
    re.I,
)

# Observability
_docx_marker_sanitized_count = 0
_docx_marker_blocked_count = 0


def _sanitize_string(value: str) -> str:
    """Remove known DOCX conversion marker tokens from a string. Normalize whitespace."""
    if not value or not isinstance(value, str):
        return value
    s = value
    for pattern, repl in _DOCX_MARKER_PATTERNS:
        s = pattern.sub(repl, s)
    # Collapse multiple spaces and strip
    s = re.sub(r"[ \t]+", " ", s).strip()
    return s


def _recursive_sanitize_resume_json(obj: Any) -> None:
    """In-place sanitize all string values in resume JSON (header, sections, entries, etc.)."""
    global _docx_marker_sanitized_count
    if isinstance(obj, dict):
        for k, v in list(obj.items()):
            if isinstance(v, str):
                cleaned = _sanitize_string(v)
                if cleaned != v:
                    _docx_marker_sanitized_count += 1
                obj[k] = cleaned
            else:
                _recursive_sanitize_resume_json(v)
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            if isinstance(item, str):
                cleaned = _sanitize_string(item)
                if cleaned != item:
                    _docx_marker_sanitized_count += 1
                obj[i] = cleaned
            else:
                _recursive_sanitize_resume_json(item)


def _scan_for_markers(obj: Any, path: str = "") -> list[tuple[str, str]]:
    """Return list of (path, value) for any string that still contains marker patterns."""
    found: list[tuple[str, str]] = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            p = f"{path}.{k}" if path else k
            if isinstance(v, str):
                if _DOCX_MARKER_ANY.search(v):
                    found.append((p, v))
            else:
                found.extend(_scan_for_markers(v, p))
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            p = f"{path}[{i}]"
            if isinstance(item, str):
                if _DOCX_MARKER_ANY.search(item):
                    found.append((p, item))
            else:
                found.extend(_scan_for_markers(item, p))
    return found


def _safety_gate_markers(resume_json: dict) -> None:
    """Ensure no formatting markers remain; sanitize once, then fail if still present."""
    global _docx_marker_blocked_count
    first_scan = _scan_for_markers(resume_json)
    if not first_scan:
        return
    _recursive_sanitize_resume_json(resume_json)
    second_scan = _scan_for_markers(resume_json)
    if second_scan:
        _docx_marker_blocked_count += 1
        sample = second_scan[0]
        raise RuntimeError(
            f"Resume JSON still contains formatting markers after sanitization (e.g. at {sample[0]}: {sample[1][:80]!r}). "
            "Conversion rejected to avoid leaking markers into PDF."
        )


# ---------------------------------------------------------------------------
# Step 2: Classify content into structured JSON via LLM
# ---------------------------------------------------------------------------

_SYSTEM_PROMPT = """\
You are a resume content classifier. Given structured text extracted from a \
.docx file, output a JSON object matching the schema below. Output ONLY valid \
JSON — no explanations, no markdown fences.

SCHEMA:
{
  "header": {
    "name": "Full Name",
    "contact_items": [{"text": "visible text", "url": "optional URL or empty"}]
  },
  "sections": [
    {
      "title": "Section Title",
      "type": "entry_list | skills | text",
      "entries": [{"primary":"Org","secondary":"Role","location":"City, ST",\
"dates":"Start - End","bullets":["..."]}],
      "items": [{"label":"Category","value":"item1, item2"}],
      "content": "freeform text"
    }
  ]
}

SECTION TYPES:
- "entry_list": Experience, Education, Projects, Awards, Certifications — \
anything with entries that have a primary title, optional subtitle/location/\
dates, and optional bullets.
- "skills": Label-value pairs (Technical Skills, Languages, etc.).
- "text": Free-form paragraph (Summary, Objective, etc.).

Include only the fields relevant to the chosen type (entries for entry_list, \
items for skills, content for text).

RULES:
- Preserve visible resume text exactly. Do not rephrase or omit content.
- Do NOT copy any helper metadata labels or tokens into output fields. Only \
the text that appears after "TEXT:" (or in table cells before alignment hints) \
is content; ALIGNMENT and STYLE_HINT lines are for your classification only.
- Bold runs usually indicate primary (company/school name). Italic usually \
secondary (role/degree). Right-aligned or after-tab content is often dates/location.
- Bullet items go in the "bullets" array.
- For contact items, infer URLs: emails -> mailto:, linkedin/github -> https://
- Join split fragments (e.g. "Jan 20" + "25" -> "Jan 2025").

FORBIDDEN in output (never put these in any JSON string value):
[BOLD], [ITALIC], [UNDERLINE], [TAB], [RIGHT-ALIGNED], [CENTER-ALIGNED], [/], \
(align: center), (align: right), (align: left), [HEADING ...], [BULLET], [TABLE ...]."""


def generate_latex_from_docx(
    structured_content: dict,
    target_pages: int = 1,
    max_retries: int = 2,
) -> tuple[str, bytes]:
    """Generate LaTeX from structured .docx content.

    Pipeline:
      1. LLM classifies extracted content into structured JSON (one call).
      2. Deterministic renderer converts JSON to LaTeX.
      3. Binary-search page fitting via local compilation (no extra LLM calls).

    Args:
        structured_content: Output from ``extract_content_from_docx()``.
        target_pages: Desired number of pages.
        max_retries: Max LLM classification attempts.

    Returns:
        Tuple of ``(latex_string, compiled_pdf_bytes)``.

    Raises:
        RuntimeError: If conversion fails after all retries.
    """
    client = OpenAI(api_key=settings.openai_api_key)

    # Use the same model configured for the rest of the app
    model = settings.model_name
    if model.startswith("openai/"):
        model = model[len("openai/"):]
    completion_token_budget = max(settings.max_tokens or 16000, 4000)
    if model.startswith("gpt-5") and completion_token_budget < 16000:
        completion_token_budget = 16000

    # Format the structured content as a readable string for the LLM
    content_description = _format_content_for_llm(structured_content)
    debug_context: dict[str, Any] = {
        "model": model,
        "target_pages": target_pages,
        "max_retries": max_retries,
        "completion_token_budget": completion_token_budget,
        "input": {
            "element_count": len(structured_content.get("elements", [])),
            "raw_text_chars": len(structured_content.get("raw_text", "") or ""),
            "formatted_prompt_chars": len(content_description),
        },
        "attempts": [],
    }

    messages: list[dict] = [
        {"role": "system", "content": _SYSTEM_PROMPT},
        {
            "role": "user",
            "content": (
                "Classify this resume content into the JSON schema:\n\n"
                + content_description
            ),
        },
    ]

    last_error = ""
    last_exception: Exception | None = None

    for attempt in range(1, max_retries + 1):
        logger.info("DOCX->JSON classification attempt %d/%d", attempt, max_retries)
        attempt_debug: dict[str, Any] = {"attempt": attempt}

        # --- Call LLM -------------------------------------------------------
        try:
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                max_completion_tokens=completion_token_budget,
                response_format={"type": "json_object"},
            )
            choice = response.choices[0]
            raw_output = choice.message.content or ""
            finish_reason = getattr(choice, "finish_reason", None)
            attempt_debug.update({
                "stage": "llm_response",
                "response_id": getattr(response, "id", None),
                "finish_reason": finish_reason,
                "usage": _usage_to_debug_dict(getattr(response, "usage", None)),
                "raw_output_chars": len(raw_output),
                "raw_output_preview": _debug_preview(raw_output),
            })
        except Exception as e:
            logger.error("OpenAI API error on attempt %d: %s", attempt, e)
            last_error = str(e)
            last_exception = e
            attempt_debug.update({
                "stage": "openai_api_error",
                "error_type": type(e).__name__,
                "error": str(e),
                "exception_traceback": traceback.format_exc(),
            })
            debug_context["attempts"].append(attempt_debug)
            continue

        # --- Parse JSON -----------------------------------------------------
        try:
            if not raw_output.strip():
                raise ValueError(f"LLM returned empty response content (finish_reason={finish_reason})")
            resume_json = _extract_json_from_response(raw_output)
            resume_json = _validate_resume_json(resume_json)
            # Post-parse sanitization: strip any marker tokens that leaked into output
            _n_sanitized_before = _docx_marker_sanitized_count
            _recursive_sanitize_resume_json(resume_json)
            if _docx_marker_sanitized_count > _n_sanitized_before:
                logger.info(
                    "DOCX marker sanitization applied (%d string(s) cleaned in this conversion)",
                    _docx_marker_sanitized_count - _n_sanitized_before,
                )
        except (ValueError, json.JSONDecodeError) as e:
            logger.warning("JSON parsing failed on attempt %d: %s", attempt, e)
            last_error = str(e)
            last_exception = e
            attempt_debug.update({
                "stage": "json_parse_error",
                "error_type": type(e).__name__,
                "error": str(e),
                "exception_traceback": traceback.format_exc(),
                "raw_output_chars": len(raw_output),
                "raw_output_preview": _debug_preview(raw_output),
            })
            debug_context["attempts"].append(attempt_debug)
            messages.append({"role": "assistant", "content": raw_output})
            messages.append({
                "role": "user",
                "content": "Invalid JSON. Output ONLY valid JSON matching the schema.",
            })
            continue

        # --- Pre-render safety gate: fail fast if markers remain -------------
        try:
            _safety_gate_markers(resume_json)
        except RuntimeError as e:
            logger.error("DOCX safety gate: %s", e)
            last_error = str(e)
            last_exception = e
            attempt_debug.update({
                "stage": "safety_gate_error",
                "error_type": type(e).__name__,
                "error": str(e),
                "exception_traceback": traceback.format_exc(),
                "raw_output_chars": len(raw_output),
                "raw_output_preview": _debug_preview(raw_output),
            })
            debug_context["attempts"].append(attempt_debug)
            messages.append({"role": "assistant", "content": raw_output})
            messages.append({
                "role": "user",
                "content": (
                    "Your output contained formatting markers that must not appear in resume text. "
                    "Output ONLY visible resume content in JSON string fields — no [BOLD], [TAB], "
                    "(align: ...), or similar tokens."
                ),
            })
            continue

        # --- Render + page-fit (deterministic, no LLM) ----------------------
        try:
            latex, pdf_bytes = fit_to_pages(resume_json, target_pages)
            logger.info("DOCX->LaTeX conversion succeeded on attempt %d", attempt)
            return latex, pdf_bytes
        except RuntimeError as e:
            logger.error("Rendering/compilation failed: %s", e)
            last_error = str(e)
            last_exception = e
            attempt_debug.update({
                "stage": "render_compile_error",
                "error_type": type(e).__name__,
                "error": str(e),
                "exception_traceback": traceback.format_exc(),
                "raw_output_chars": len(raw_output),
                "raw_output_preview": _debug_preview(raw_output),
            })
            debug_context["attempts"].append(attempt_debug)
            raise DOCXConversionError(
                f"DOCX to LaTeX rendering failed. Last error: {last_error}",
                debug_context,
            ) from e

    debug_context["last_error"] = last_error
    raise DOCXConversionError(
        f"DOCX to LaTeX conversion failed after {max_retries} attempts. "
        f"Last error: {last_error}",
        debug_context,
    ) from last_exception


def _format_content_for_llm(structured_content: dict) -> str:
    """Format extracted .docx content for the LLM with visible text only; metadata separate.

    Emits TEXT: <visible content> and ALIGNMENT/STYLE_HINT on separate lines so the LLM
    never sees bracket-style markers in the content it should copy into JSON.
    """
    lines: list[str] = []
    lines.append("=== DOCUMENT CONTENT ===\n")

    for elem in structured_content.get("elements", []):
        if elem["type"] == "paragraph":
            # Build visible text only (no [BOLD], [TAB], etc.)
            visible_parts: list[str] = []
            tab_align = elem.get("tab_alignment")
            for run in elem.get("runs", []):
                if run.get("tab"):
                    # Tab: emit a single space as separator; alignment is metadata below
                    visible_parts.append(" ")
                    continue
                visible_parts.append(run.get("text", ""))
            text_only = "".join(visible_parts).strip()

            # Metadata on separate lines so LLM does not copy them as content
            if elem.get("is_heading"):
                lines.append(f"STYLE_HINT: heading level={elem.get('heading_level', 1)}")
            elif elem.get("is_list"):
                lines.append("STYLE_HINT: bullet")
            if elem.get("alignment") != "left":
                lines.append(f"ALIGNMENT: {elem['alignment']}")
            if tab_align:
                lines.append(f"TAB_ALIGNMENT: {tab_align} (content after tab is often dates/location)")
            # Run formatting hint for classification only (bold=primary, italic=secondary)
            run_hints = []
            for run in elem.get("runs", []):
                if run.get("tab"):
                    continue
                if run.get("bold"):
                    run_hints.append("bold")
                elif run.get("italic"):
                    run_hints.append("italic")
                elif run.get("underline"):
                    run_hints.append("underline")
            if run_hints:
                lines.append(f"STYLE_HINT: runs have {', '.join(run_hints)} (for primary/secondary classification)")

            lines.append(f"TEXT: {text_only}")

        elif elem["type"] == "table":
            lines.append(f"TABLE: {elem['num_rows']}x{elem['num_cols']}")
            for row in elem.get("rows", []):
                cell_texts = [cell.get("text", "") for cell in row]
                lines.append(" | ".join(cell_texts))
                # Alignment as metadata only (one line per row)
                alignments = [cell.get("alignment", "left") for cell in row]
                if any(a != "left" for a in alignments):
                    lines.append(f"ALIGNMENT: {', '.join(alignments)}")
            lines.append("")

    return "\n".join(lines)


def _extract_json_from_response(raw: str) -> dict:
    """Extract and parse JSON from the LLM response."""
    raw = raw.strip()

    # Strip markdown fences if present
    if raw.startswith("```"):
        first_newline = raw.index("\n") if "\n" in raw else len(raw)
        raw = raw[first_newline + 1:]
        if raw.endswith("```"):
            raw = raw[:-3]
        raw = raw.strip()

    # Try direct parse
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    # Try to locate a JSON object in the response
    start = raw.find("{")
    end = raw.rfind("}") + 1
    if start >= 0 and end > start:
        try:
            return json.loads(raw[start:end])
        except json.JSONDecodeError:
            pass

    raise ValueError(f"Could not parse JSON from LLM response: {raw[:200]}")


def _validate_resume_json(data: dict) -> dict:
    """Validate and normalise the resume JSON to match the expected schema."""
    if not isinstance(data, dict):
        raise ValueError("Resume JSON must be a dict")

    # Header
    header = data.get("header", {})
    if not isinstance(header, dict):
        header = {}
    header.setdefault("name", "")
    header.setdefault("contact_items", [])
    data["header"] = header

    # Sections
    sections = data.get("sections", [])
    if not isinstance(sections, list):
        sections = []

    validated: list[dict] = []
    for section in sections:
        if not isinstance(section, dict):
            continue
        stype = section.get("type", "entry_list")
        if stype not in ("entry_list", "skills", "text"):
            stype = "entry_list"
        section["type"] = stype
        section.setdefault("title", "")

        if stype == "entry_list":
            entries = section.get("entries", [])
            if not isinstance(entries, list):
                entries = []
            for entry in entries:
                if not isinstance(entry, dict):
                    continue
                entry.setdefault("primary", "")
                entry.setdefault("secondary", "")
                entry.setdefault("location", "")
                entry.setdefault("dates", "")
                entry.setdefault("bullets", [])
            section["entries"] = [e for e in entries if isinstance(e, dict)]
        elif stype == "skills":
            section.setdefault("items", [])
        elif stype == "text":
            section.setdefault("content", "")

        validated.append(section)

    data["sections"] = validated
    return data
